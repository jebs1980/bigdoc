"""
Génère un document Word de relecture des bilans eval.
Usage : python3 generate_eval_word.py [eval_results.json] [output.docx]
"""
import json, sys, os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS_FILE = Path(sys.argv[1] if len(sys.argv) > 1 else "eval_results.json")
OUTPUT_FILE  = Path(sys.argv[2] if len(sys.argv) > 2 else "eval_relecture.docx")

SCORE_COLORS = {
    "critique":   RGBColor(0xDC, 0x26, 0x26),
    "fragile":    RGBColor(0xEA, 0x58, 0x0C),
    "passable":   RGBColor(0xCA, 0x8A, 0x04),
    "bon":        RGBColor(0x16, 0xA3, 0x4A),
}

DIM_LABELS = {
    "administration":  "Administration & CPAM",
    "achats_materiel": "Achats & Matériel",
    "informatique":    "Informatique & Infrastructure",
    "comptabilite":    "Comptabilité & Finances",
    "charge_mentale":  "Charge mentale",
    "financement":     "Financement",
    "developpement":   "Développement",
}

def score_color(s):
    if s is None: return RGBColor(0x6B, 0x72, 0x80)
    if s <= 25:   return SCORE_COLORS["critique"]
    if s <= 50:   return SCORE_COLORS["fragile"]
    if s <= 75:   return SCORE_COLORS["passable"]
    return SCORE_COLORS["bon"]

def add_shaded_para(doc, text, bg_hex="EEF2FF", font_size=10, bold=False, color=None):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_hex)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    return p

def generate(results):
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Page de titre
    title = doc.add_heading("Bigdoc — Relecture des bilans", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Évaluation du prompt · {datetime.now().strftime('%d/%m/%Y')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    sub.runs[0].font.size = Pt(11)

    ok_count = sum(1 for r in results if r.get("status") == "ok")
    stats = doc.add_paragraph(f"{ok_count} bilans générés sur {len(results)} cas de test")
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats.runs[0].font.size = Pt(10)
    stats.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    doc.add_page_break()

    # Table des matières simplifiée
    doc.add_heading("Index des cas", 1)
    for r in results:
        b = r.get("bilan", {})
        score = b.get("score_global")
        p = doc.add_paragraph(style="List Bullet")
        run_id = p.add_run(f"{r['id']} — ")
        run_id.font.bold = True
        run_id.font.size = Pt(10)
        run_label = p.add_run(f"{r.get('specialite','?')} · {r.get('ville','?').split('(')[0].strip()}")
        run_label.font.size = Pt(10)
        if score is not None:
            run_score = p.add_run(f" [{score}/100]")
            run_score.font.bold = True
            run_score.font.size = Pt(10)
            run_score.font.color.rgb = score_color(score)

    doc.add_page_break()

    # Bilans
    for r in results:
        b = r.get("bilan", {})
        if r.get("status") == "erreur":
            doc.add_heading(f"{r['id']} — ERREUR", 2)
            doc.add_paragraph(f"Spécialité : {r.get('specialite','?')} · Ville : {r.get('ville','?')}")
            doc.add_paragraph("Ce cas n'a pas pu être traité.").runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            doc.add_page_break()
            continue

        score = b.get("score_global")
        col = score_color(score)

        # Titre du cas
        h = doc.add_heading(f"{r['id']} — {r.get('label', '')}", 2)

        # Bandeau score
        info = doc.add_paragraph()
        s1 = info.add_run(f"Score : {score}/100  ")
        s1.font.bold = True
        s1.font.size = Pt(13)
        s1.font.color.rgb = col
        s2 = info.add_run(f"{b.get('titre_personnalise', b.get('niveau',''))}")
        s2.font.size = Pt(11)
        s2.font.bold = True
        info.paragraph_format.space_after = Pt(4)

        # Spécialité / Ville / Phase
        meta = doc.add_paragraph()
        meta.add_run(f"Spécialité : {r.get('specialite','?')}  ·  Ville : {r.get('ville','?')}  ·  Phase : {b.get('phase','?')}")
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Ce que le médecin a écrit
        texte_libre = r.get("texte_libre", "").strip()
        if texte_libre:
            doc.add_paragraph("Ce que le médecin a écrit :").runs[0].font.bold = True
            add_shaded_para(doc, texte_libre, bg_hex="FFF7ED", font_size=9)

        # Insight
        insight = b.get("ce_que_vous_nignorez_probablement_pas", "")
        if insight:
            doc.add_paragraph()
            ip = doc.add_paragraph()
            ir = ip.add_run(f"💡 {insight}")
            ir.font.italic = True
            ir.font.size = Pt(10)
            ir.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

        # Message bilan
        doc.add_paragraph("Message bilan :").runs[0].font.bold = True
        add_shaded_para(doc, b.get("message_bilan", "—"), bg_hex="F5F3FF", font_size=10)

        # Alerte urgente
        alerte = b.get("alerte_urgente", "").strip()
        if alerte:
            add_shaded_para(doc, f"⚠️ {alerte}", bg_hex="FEF2F2", font_size=9, color=RGBColor(0x99, 0x1B, 0x1B))

        # Dimensions
        doc.add_paragraph("Dimensions :").runs[0].font.bold = True
        dims = b.get("dimensions", {})
        for key, d in dims.items():
            dp = doc.add_paragraph()
            dr1 = dp.add_run(f"{DIM_LABELS.get(key, key)} : {d.get('score','?')}/20  ")
            dr1.font.bold = True
            dr1.font.size = Pt(9)
            dr1.font.color.rgb = score_color(int(d.get('score', 10)) * 5)
            dr2 = dp.add_run(d.get("commentaire", "")[:200])
            dr2.font.size = Pt(9)
            dp.paragraph_format.space_before = Pt(1)
            dp.paragraph_format.space_after = Pt(1)

        # Quick wins
        qws = b.get("quick_wins", [])
        if qws:
            doc.add_paragraph("Actions RMS proposées :").runs[0].font.bold = True
            for i, qw in enumerate(qws, 1):
                qp = doc.add_paragraph(style="List Number")
                qr = qp.add_run(qw)
                qr.font.size = Pt(9)

        # Recommandation principale
        reco = b.get("recommandation_principale", {})
        if reco:
            doc.add_paragraph("Recommandation :").runs[0].font.bold = True
            rp = doc.add_paragraph()
            rp.add_run(f"{reco.get('service','?')} — {reco.get('tarif','?')}").font.bold = True
            rp.runs[0].font.size = Pt(10)
            rp.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            add_shaded_para(doc, reco.get("justification", ""), bg_hex="EEF2FF", font_size=9)

        # Zone de commentaire pour la relectrice
        doc.add_paragraph()
        zc_title = doc.add_paragraph("📝 Zone de commentaire :")
        zc_title.runs[0].font.bold = True
        zc_title.runs[0].font.size = Pt(9)
        zc_title.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        for _ in range(3):
            zc = doc.add_paragraph("_" * 80)
            zc.runs[0].font.size = Pt(9)
            zc.runs[0].font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
            zc.paragraph_format.space_before = Pt(1)
            zc.paragraph_format.space_after = Pt(1)

        doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print(f"✅ {OUTPUT_FILE} — {len(results)} cas")

if __name__ == "__main__":
    if not RESULTS_FILE.exists():
        print(f"❌ {RESULTS_FILE} introuvable")
        sys.exit(1)
    results = json.loads(RESULTS_FILE.read_text())
    results.sort(key=lambda r: r.get("id",""))
    generate(results)
